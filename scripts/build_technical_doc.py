"""Build the RentalFlow Technical Reference Word document.

Target reader: a developer who is about to start writing code in this repo.
Structure: welcome + local setup + workflow + recipes, then deep architecture
chapters reused from the audit (tech stack, multi-tenant, routing, schema, API,
crons, integrations, security, ops, debt).

Output: ./RentalFlow_Technical_Reference.docx
"""

import os
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
sys.path.insert(0, os.path.join(REPO, 'audit-data'))

import _renderer as R


def safe_import(name):
    try:
        return __import__(name)
    except Exception as e:
        print(f'  [skip] {name}: {e}')
        return None


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RENTALFLOW')
    r.bold = True
    r.font.size = Pt(56)
    r.font.color.rgb = R.NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SaaS for Rental Businesses')
    r.font.size = Pt(13)
    r.font.color.rgb = R.GOLD
    r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Technical Reference\nfor New Engineers')
    r.font.size = Pt(28)
    r.font.color.rgb = R.INK
    r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Everything you need to start writing code in this repo.")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = R.GRAY

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Prepared ' + date.today().strftime('%B %d, %Y'))
    r.font.size = Pt(11)
    r.font.color.rgb = R.GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Use side-by-side with the codebase. This document tells you what '
                  'to read and where to start.')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = R.GRAY_L

    doc.add_page_break()


def main():
    doc = Document()
    R.set_styles(doc)
    add_cover(doc)

    toc_entries = [
        (1,  'Welcome',                     'A 10-minute orientation. Read this first.'),
        (2,  'Local Setup',                 'Clone → run dev → first commit. ~30 minutes.'),
        (3,  'Daily Workflow',              'Git, deploys, testing, hotfixes. The rules of the road.'),
        (4,  'Common Recipes',              'How to add a page, an email template, an integration.'),
        (5,  'Technology Stack',            'Every framework, every service, why it was chosen.'),
        (6,  'Multi-Tenant Architecture',   'The most important design topic. Read carefully.'),
        (7,  'Routing & Middleware',        'How a request finds its tenant in 30 ms.'),
        (8,  'Booking Wizard Deep Dive',    'The hottest server path — pricing, holds, payment.'),
        (9,  'Admin Panel — 81 Surfaces',   'Per-page reference (data sheets from the JSON inventory).'),
        (10, 'API Reference',               'All 51 endpoints — contracts, auth, rate limits.'),
        (11, 'Background Jobs',             '13 crons orchestrating the lifecycle.'),
        (12, 'Data Dictionary',             '88 tables, 197 RLS policies, 41 functions.'),
        (13, 'Integrations',                '11 services — Stripe, GHL, Twilio, Resend, IMAP, AI.'),
        (14, 'Security Architecture',       'Authn, authz, secret management, RLS.'),
        (15, 'Operations & Deploy',         'Vercel, Sentry, restore drill, incidents.'),
        (16, 'Tech Debt by the Numbers',    'TODOs, any-usage, test coverage — where it actually lives.'),
    ]
    R.add_toc_simple(doc, toc_entries,
        header='Contents',
        subhead='Read 1-4 in order. After that, jump to whatever you need.')

    api = R.make_api(doc)

    chapters = [
        ('tech_01_welcome',                'render'),
        ('tech_02_local_setup',            'render'),
        ('tech_03_workflow',               'render'),
        ('tech_04_recipes',                'render'),
        ('chapter_03_tech_stack',          'render'),
        ('chapter_04_multitenant',         'render'),
        ('chapter_05_routing',             'render'),
        ('chapter_09_booking_wizard',      'render'),
        ('chapter_10_admin',               'render'),
        ('chapter_12_api',                 'render'),
        ('chapter_13_crons',               'render'),
        ('chapter_14_schema',              'render'),
        ('chapter_15_integrations',        'render'),
        ('chapter_16_security',            'render'),
        ('chapter_17_operations',          'render'),
        ('chapter_18_tech_debt',           'render'),
    ]

    for mod_name, fn_name in chapters:
        mod = safe_import(mod_name)
        if mod and hasattr(mod, fn_name):
            print(f'  rendering {mod_name}...')
            getattr(mod, fn_name)(api)

    out_path = os.path.join(REPO, 'RentalFlow_Technical_Reference.docx')
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f'\nSaved: {out_path} ({size_kb} KB)')


if __name__ == '__main__':
    main()
