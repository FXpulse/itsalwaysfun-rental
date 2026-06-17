"""Build the RentalFlow Client Guide Word document.

Target reader: owner of a rental business (inflatables, party rentals, equipment).
Tone: owner-to-owner. Practical. No tech jargon. Benefit-first.
Length: ~25-40 pages.

Output: ./RentalFlow_For_Rental_Owners.docx
"""

import os
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
sys.path.insert(0, os.path.join(REPO, 'audit-data'))

import _renderer as R


def safe_import(name):
    try:
        return __import__(name)
    except Exception as e:
        print(f'  [skip] {name}: {e}')
        return None


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RENTALFLOW')
    r.bold = True
    r.font.size = Pt(56)
    r.font.color.rgb = R.NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Built by a rental business.\nFor rental businesses.')
    r.font.size = Pt(13)
    r.font.color.rgb = R.GOLD
    r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Run Your Whole Business\non One Platform')
    r.font.size = Pt(30)
    r.font.color.rgb = R.INK
    r.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bookings, payments, dispatch, customers, marketing.\nAll for $99 a month, flat.")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = R.GRAY

    for _ in range(6):
        doc.add_paragraph()

    stats = [
        ('$99', 'flat /mo',     R.NAVY),
        ('0%',  'per-booking',  R.GREEN),
        ('14',  'day trial',    R.BLUE),
        ('LIVE','in production', R.TEAL),
    ]
    R.add_stats_strip(doc, stats)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('A guide for rental business owners.')
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = R.GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('getrentalflow.com')
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = R.NAVY

    doc.add_page_break()


def main():
    doc = Document()
    R.set_styles(doc)
    add_cover(doc)

    toc_entries = [
        (1, 'Why RentalFlow Exists',         'Built by an owner who got tired of the existing tools.'),
        (2, 'Everything You Get',            'One page. The whole platform at a glance.'),
        (3, 'Your Customers\' Experience',   'What customers see — from browse to book to come back.'),
        (4, 'Your Office',                   'Bookings, calendar, dispatch, drivers, accounting.'),
        (5, 'How You Get Paid',              'Stripe Connect. Your money flows to your bank in days.'),
        (6, 'Grow Your Business',            'Loyalty, referrals, reviews, email & SMS, all built in.'),
        (7, 'What You Can Stop Paying For',  'The 6 tools you probably have today that we replace.'),
        (8, 'Pricing — and What It Means',   "$99 flat. No per-booking. No upsells. No surprises."),
        (9, 'Getting Started',               'What the first 14 days look like.'),
    ]
    R.add_toc_simple(doc, toc_entries,
        header='What\'s Inside',
        subhead='9 short chapters · 30-minute read · written by someone who runs a rental business.')

    api = R.make_api(doc)

    chapters = [
        ('client_01_intro',          'render'),
        ('client_02_what_you_get',   'render'),
        ('client_03_customer_side',  'render'),
        ('client_04_office',         'render'),
        ('client_05_payments',       'render'),
        ('client_06_marketing',      'render'),
        ('client_07_replaces',       'render'),
        ('client_08_pricing',        'render'),
        ('client_09_get_started',    'render'),
    ]

    for mod_name, fn_name in chapters:
        mod = safe_import(mod_name)
        if mod and hasattr(mod, fn_name):
            print(f'  rendering {mod_name}...')
            getattr(mod, fn_name)(api)

    out_path = os.path.join(REPO, 'RentalFlow_For_Rental_Owners.docx')
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f'\nSaved: {out_path} ({size_kb} KB)')


if __name__ == '__main__':
    main()
