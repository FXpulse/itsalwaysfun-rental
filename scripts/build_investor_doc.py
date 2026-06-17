"""Build the RentalFlow Investor Brief Word document.

Target reader: angel / pre-seed investor evaluating a $50-500K check.
Tone: founder-direct, no TAM-bullshit, technical defensibility is real.
Length: ~40-60 pages, polished.

Output: ./RentalFlow_Investor_Brief.docx
"""

import os
import sys
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
sys.path.insert(0, os.path.join(REPO, 'audit-data'))

import _renderer as R


def safe_import(name):
    try:
        return __import__(name)
    except Exception as e:
        print(f'  [skip] {name}: {e}')
        return None


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RENTALFLOW')
    r.bold = True
    r.font.size = Pt(56)
    r.font.color.rgb = R.NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Vertical SaaS for Rental Businesses')
    r.font.size = Pt(13)
    r.font.color.rgb = R.GOLD
    r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Investor Brief')
    r.font.size = Pt(32)
    r.font.color.rgb = R.INK
    r.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bootstrapped to LIVE. Looking for the right angel partner.")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = R.GRAY

    for _ in range(6):
        doc.add_paragraph()

    # Quick stats strip on cover
    stats = [
        ('LIVE', 'production', R.GREEN),
        ('$99', 'flat /mo',   R.NAVY),
        ('11',   'integrations', R.BLUE),
        ('0',    'rounds raised', R.PURPLE),
    ]
    R.add_stats_strip(doc, stats)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Prepared ' + date.today().strftime('%B %d, %Y'))
    r.font.size = Pt(11)
    r.font.color.rgb = R.GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('itsalwaysfun.com  ·  getrentalflow.com')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = R.GRAY_L

    doc.add_page_break()


def main():
    doc = Document()
    R.set_styles(doc)
    add_cover(doc)

    toc_entries = [
        (1, 'The One-Pager',             'What RentalFlow is. Why now. What the check buys.'),
        (2, 'The Market',                'Rental software is overpriced, underpowered, and fragmented.'),
        (3, 'The Product',               'Six surfaces, one platform. 11 integrations live.'),
        (4, 'The Moat',                  'Why a copycat would take 18 months to catch up.'),
        (5, 'Traction',                  'IAF is the flagship and a paying customer. Early signal.'),
        (6, 'Unit Economics',            'Why a $99 flat plan is the right shape at this stage.'),
        (7, 'The Team',                  "How Ludmila + Henry ship at this pace."),
        (8, 'Risks',                     'Honest assessment. The 5 things that could go wrong.'),
        (9, 'The Ask',                   'What we want. What you get. What happens next.'),
    ]
    R.add_toc_simple(doc, toc_entries,
        header='Contents',
        subhead='9 chapters · 40-minute read · founder-direct.')

    api = R.make_api(doc)

    chapters = [
        ('inv_01_one_pager',  'render'),
        ('inv_02_market',     'render'),
        ('inv_03_product',    'render'),
        ('inv_04_moat',       'render'),
        ('inv_05_traction',   'render'),
        ('inv_06_economics',  'render'),
        ('inv_07_team',       'render'),
        ('inv_08_risks',      'render'),
        ('inv_09_ask',        'render'),
    ]

    for mod_name, fn_name in chapters:
        mod = safe_import(mod_name)
        if mod and hasattr(mod, fn_name):
            print(f'  rendering {mod_name}...')
            getattr(mod, fn_name)(api)

    out_path = os.path.join(REPO, 'RentalFlow_Investor_Brief.docx')
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f'\nSaved: {out_path} ({size_kb} KB)')


if __name__ == '__main__':
    main()
