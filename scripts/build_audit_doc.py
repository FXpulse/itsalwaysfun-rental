"""Build the RentalFlow Technical Audit Word document.

Goal: presentation-quality, ~150-200 page deep audit of the system.
Sections work both standalone (for individual audiences) and together
(as a single comprehensive deliverable).

Run: python scripts/build_audit_doc.py
Output: ./RentalFlow_Technical_Audit.docx

Content modules live in audit-data/ (one .py module per topic) so this
file owns rendering, not content. That keeps the rendering code legible.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os
import sys

# Allow audit-data/ as a content module location
HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
sys.path.insert(0, os.path.join(REPO, "audit-data"))

# ─── Palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x1A, 0x6E)
NAVY_D = RGBColor(0x0E, 0x0E, 0x3F)  # darker navy for chapter bands
GOLD   = RGBColor(0xFF, 0xC9, 0x29)
INK    = RGBColor(0x1F, 0x29, 0x37)
GRAY   = RGBColor(0x4B, 0x5B, 0x68)
GRAY_L = RGBColor(0x8B, 0x96, 0xA3)
PAPER  = RGBColor(0xF8, 0xFA, 0xFC)
GREEN  = RGBColor(0x0E, 0x7C, 0x4A)
AMBER  = RGBColor(0xB4, 0x73, 0x0C)
RED    = RGBColor(0xB3, 0x26, 0x1A)
BLUE   = RGBColor(0x1B, 0x6A, 0xBC)
TEAL   = RGBColor(0x0F, 0x76, 0x82)
PURPLE = RGBColor(0x6B, 0x46, 0xC1)


# ─── Low-level XML helpers ────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Paint a cell with a solid fill color (hex string, no #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CFD8DC", size="4", sides=("top", "bottom", "left", "right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def set_paragraph_border(paragraph, color="1A1A6E", size="6", sides=("left",)):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '6')
        b.set(qn('w:color'), color)
        borders.append(b)
    p_pr.append(borders)


def rgb_to_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# ─── Style setup ──────────────────────────────────────────────────────

def set_styles(doc):
    """Configure document-wide typography."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.18

    for h, size, color in (
        ('Heading 1', 26, NAVY),
        ('Heading 2', 16, NAVY),
        ('Heading 3', 12, NAVY),
        ('Heading 4', 11, GRAY),
    ):
        s = doc.styles[h]
        s.font.name = 'Calibri'
        s.font.color.rgb = color
        s.font.size = Pt(size)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(4)

    # Make margins a bit tighter so we get more content per page
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)


# ─── High-level building blocks ───────────────────────────────────────

def add_p(doc, text, *, bold=False, italic=False, color=None, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color is not None: r.font.color.rgb = color
    if size is not None: r.font.size = Pt(size)
    return p


def add_runs(doc, runs, *, align=None):
    """Add a paragraph with mixed-style runs.
    runs = list of (text, dict_of_style) where style keys: bold, italic, color, size."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for text, style in runs:
        r = p.add_run(text)
        r.bold = style.get('bold', False)
        r.italic = style.get('italic', False)
        if 'color' in style: r.font.color.rgb = style['color']
        if 'size' in style: r.font.size = Pt(style['size'])
    return p


def add_h1(doc, text):
    return doc.add_heading(text, level=1)

def add_h2(doc, text):
    return doc.add_heading(text, level=2)

def add_h3(doc, text):
    return doc.add_heading(text, level=3)

def add_h4(doc, text):
    return doc.add_heading(text, level=4)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))
    if isinstance(text, list):
        for chunk, is_bold in text:
            r = p.add_run(chunk)
            r.bold = is_bold
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))
    p.add_run(text)
    return p


# ─── Visual elements: chapter divider, callouts, cards, badges ────────

def add_chapter_divider(doc, number, title, tagline, audience_tags=None):
    """A big visual chapter title — meant to occupy ~3/4 page.

    Format:
      [Tall navy band with white number on left, title + tagline on right]
      [audience badges underneath]
    """
    doc.add_page_break()

    # Top spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(80)

    # Big chapter number in gold
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"CHAPTER {number:02d}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(2)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_border(p, color=rgb_to_hex(NAVY), size="18", sides=("bottom",))

    # Tagline
    p = doc.add_paragraph()
    r = p.add_run(tagline)
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(10)

    # Audience tags
    if audience_tags:
        tbl = doc.add_table(rows=1, cols=len(audience_tags))
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, tag in enumerate(audience_tags):
            cell = tbl.rows[0].cells[i]
            cell.text = ''
            set_cell_shading(cell, rgb_to_hex(NAVY))
            cell.width = Inches(1.5)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('  ' + tag.upper() + '  ')
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Bottom spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(40)


def add_section_break(doc):
    """A subtle visual rule between sub-sections."""
    p = doc.add_paragraph()
    set_paragraph_border(p, color="CFD8DC", size="6", sides=("bottom",))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def add_callout(doc, kind, text, label=None):
    """A visually distinct callout box (info / warn / risk / good / quote).

    Implemented as a 1-cell table with a colored left border + tinted fill.
    """
    config = {
        'info':  ('FYI',     NAVY,   'E3E6F4'),
        'warn':  ('HEADS UP', AMBER,  'FFF4D6'),
        'risk':  ('RISK',    RED,    'F9DDD6'),
        'good':  ('STRONG',  GREEN,  'D9F2E2'),
        'quote': ('PULLOUT', PURPLE, 'EDE6F8'),
        'fact':  ('AT A GLANCE', TEAL, 'D9F0F2'),
    }
    label_default, color, fill_hex = config.get(kind, ('NOTE', GRAY, 'F0F0F0'))
    final_label = (label or label_default).upper()

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, fill_hex)
    set_cell_borders(cell, color=rgb_to_hex(color), size="18", sides=("left",))
    set_cell_borders(cell, color="DDE3E8", size="2", sides=("top", "bottom", "right"))
    cell.width = Inches(7.1)

    # Label
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(final_label)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = color

    # Body
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK

    # Spacer below
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_kv_table(doc, headers, rows, col_widths=None, header_fill=NAVY):
    """Standard tabular layout with navy header band."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = ''
        set_cell_shading(cell, rgb_to_hex(header_fill))
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
    for row in rows:
        row_cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cell = row_cells[i]
            cell.text = ''
            set_cell_borders(cell, color="E0E4E8", size="4")
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl


def add_stats_strip(doc, stats):
    """Row of mini-stat cards. stats = [(big_number, label, color), ...]"""
    tbl = doc.add_table(rows=1, cols=len(stats))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (number, label, color) in enumerate(stats):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        set_cell_shading(cell, 'F4F6F8')
        set_cell_borders(cell, color="E0E4E8", size="2")
        cell.width = Inches(1.5)

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        r = p1.add_run(str(number))
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = color

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = GRAY
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_data_sheet(doc, *, title, subtitle=None, fields):
    """Per-page / per-feature data sheet.

    fields = list of (label, body) tuples. body can be str or list[str].
    Renders as a 2-column table with label in NAVY caps + body in INK.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_border(p, color=rgb_to_hex(NAVY), size="14", sides=("left",))
    r = p.add_run(' ' + title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.1)
        r = p2.add_run(subtitle)
        r.italic = True
        r.font.color.rgb = GRAY
        r.font.size = Pt(9.5)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    for label, body in fields:
        row = tbl.add_row().cells
        # Label cell
        label_cell = row[0]
        label_cell.width = Inches(1.4)
        label_cell.text = ''
        set_cell_shading(label_cell, 'F4F6F8')
        set_cell_borders(label_cell, color="E0E4E8", size="2")
        pl = label_cell.paragraphs[0]
        rl = pl.add_run(label.upper())
        rl.bold = True
        rl.font.size = Pt(8.5)
        rl.font.color.rgb = NAVY
        # Body cell
        body_cell = row[1]
        body_cell.width = Inches(5.7)
        body_cell.text = ''
        set_cell_borders(body_cell, color="E0E4E8", size="2")
        if isinstance(body, list):
            for j, line in enumerate(body):
                if j == 0:
                    bp = body_cell.paragraphs[0]
                else:
                    bp = body_cell.add_paragraph()
                bp.paragraph_format.space_after = Pt(2)
                rb = bp.add_run('• ' + line)
                rb.font.size = Pt(9.5)
        else:
            bp = body_cell.paragraphs[0]
            rb = bp.add_run(str(body))
            rb.font.size = Pt(9.5)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def add_mono_block(doc, text, *, title=None):
    """ASCII / monospace block — used for sequence diagrams + code snippets."""
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = NAVY
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.1)
    cell.text = ''
    set_cell_shading(cell, '0E0E3F')
    set_cell_borders(cell, color="0E0E3F", size="6")
    for i, line in enumerate(text.split('\n')):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(8.8)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_pill_row(doc, pills):
    """A row of inline pill-badges. pills = [(text, color), ...]"""
    tbl = doc.add_table(rows=1, cols=len(pills))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (text, color) in enumerate(pills):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        set_cell_shading(cell, rgb_to_hex(color))
        cell.width = Inches(1.0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_2col(doc, left_text, right_text, *, left_title=None, right_title=None):
    """Two-column comparison / parallel content."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    left_cell, right_cell = tbl.rows[0].cells
    left_cell.width = Inches(3.5)
    right_cell.width = Inches(3.5)
    set_cell_borders(left_cell, color="E0E4E8", size="2")
    set_cell_borders(right_cell, color="E0E4E8", size="2")

    def fill(cell, title, body):
        cell.text = ''
        if title:
            p = cell.paragraphs[0]
            r = p.add_run(title.upper())
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = NAVY
            cell.add_paragraph()
        p2 = cell.add_paragraph()
        r = p2.add_run(body)
        r.font.size = Pt(9.5)

    fill(left_cell, left_title, left_text)
    fill(right_cell, right_title, right_text)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# ─── Cover + back-cover ──────────────────────────────────────────────

def add_cover(doc):
    # Tall top spacer (3-inch margin from top)
    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # "RENTALFLOW" big
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RENTALFLOW')
    r.bold = True
    r.font.size = Pt(56)
    r.font.color.rgb = NAVY

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SaaS for Rental Businesses')
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD
    r.bold = True

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(40)

    # Doc title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Technical Audit\n& Architecture Review')
    r.font.size = Pt(28)
    r.font.color.rgb = INK
    r.bold = True

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(30)

    # Flagship pitch
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Multi-tenant platform powering ')
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    r2 = p.add_run("It's Always Fun, LLC")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = INK
    r3 = p.add_run(' and ready for the next 10 tenants')
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(80)

    # Audience matrix
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('PREPARED FOR')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    audiences = [
        ('OWNER', 'Ludmila', NAVY),
        ('OPERATIONS', 'IAF team', TEAL),
        ('ENGINEERING', 'New hires', BLUE),
        ('M&A', 'Acquirers', PURPLE),
    ]
    for i, (heading, sub, color) in enumerate(audiences):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(1.5)
        cell.text = ''
        set_cell_shading(cell, rgb_to_hex(color))
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(heading)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(sub)
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(60)

    # Date footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Prepared ' + date.today().strftime('%B %d, %Y'))
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('itsalwaysfun.com  ·  getrentalflow.com')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY_L

    doc.add_page_break()


def add_how_to_read(doc):
    add_h1(doc, 'How to Read This Document')

    add_p(doc, 'This audit is structured to work for four different readers. Each '
                'chapter is self-contained; you do not need to read in order.')

    add_h3(doc, 'Reading paths by audience')
    add_kv_table(doc,
        ['You are…', 'Read these chapters', 'Skim the rest'],
        [
            ('Ludmila (operator)',
             '1, 2, 5, 6, 7, 8, 9 (booking), 10 (admin tour), 17 (ops)',
             '12 (API), 14 (schema) — skim'),
            ('A new engineer',
             'All — start with 3, 4, then 14, 12, 10',
             'Read 6, 7, 8 last'),
            ('A potential acquirer',
             '1, 2, 4, 13, 15, 16, 18, 19',
             '6-12 for surface inventory'),
            ('An investor / due diligence',
             '1, 2, 13 (integrations), 15 (security), 18 (debt), 19 (roadmap)',
             '14 (schema) for credibility check'),
        ],
        col_widths=[1.6, 3.3, 2.2])

    add_h3(doc, 'Visual conventions')
    add_kv_table(doc,
        ['Element', 'Meaning'],
        [
            ('Navy band callout', 'Informational note worth pausing on'),
            ('Green callout', 'A strength of the system'),
            ('Amber callout', 'Heads-up — risk or known gap'),
            ('Red callout', 'Live risk that needs action'),
            ('Teal "At a glance"', 'Headline stat for the chapter'),
            ('Monospace black block', 'Sequence diagram, flow, or code excerpt'),
            ('Two-column data sheet', 'A page/endpoint/table profile'),
        ],
        col_widths=[2.0, 5.1])
    doc.add_page_break()


def add_toc(doc, entries):
    add_h1(doc, 'Contents')
    add_p(doc, 'A 19-chapter walkthrough of the platform from cover to schema.',
          italic=True, color=GRAY)

    add_section_break(doc)

    for num, title, tagline in entries:
        # Chapter number + title row
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f'{num:02d}   ')
        r.font.size = Pt(11)
        r.font.color.rgb = GOLD
        r.bold = True
        r2 = p.add_run(title)
        r2.font.size = Pt(11)
        r2.font.color.rgb = INK
        r2.bold = True

        # Tagline
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.55)
        p2.paragraph_format.space_after = Pt(6)
        r = p2.add_run(tagline)
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = GRAY

    doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────
# Content modules — each is a function that takes the doc and adds
# its chapter. Loaded from audit-data/ where possible.
# ──────────────────────────────────────────────────────────────────────

# Chapter content lives in subsequent files attached to this builder.
# We import them lazily so a missing module doesn't break the whole run.

def safe_import(name):
    try:
        return __import__(name)
    except Exception as e:
        print(f'  [skip] could not import {name}: {e}')
        return None


def main():
    doc = Document()
    set_styles(doc)

    # Cover + intro pages
    add_cover(doc)
    add_how_to_read(doc)

    toc_entries = [
        (1, 'Executive Summary', 'The 5-minute version: what RentalFlow is and what it is not.'),
        (2, 'System at a Glance', 'Numbers, surfaces, integrations — the dashboard view.'),
        (3, 'Technology Stack', 'Every framework, every service, why it was chosen.'),
        (4, 'Multi-Tenant Architecture', 'How one codebase serves 10+ businesses safely.'),
        (5, 'Routing & Middleware', 'How a single request finds its tenant in 30ms.'),
        (6, 'Public Tenant Site', 'Every page a paying customer sees on itsalwaysfun.com.'),
        (7, 'Customer Portal', 'The logged-in customer experience — loyalty, referrals, account.'),
        (8, 'Driver Mobile App', 'PWA for the team on the road.'),
        (9, 'Booking Wizard Deep Dive', 'The hottest path — from date picker to confirmation.'),
        (10, 'Admin Panel — 56 Surfaces', 'The operator console, page-by-page data sheets.'),
        (11, 'Superadmin Platform', 'The SaaS owner backoffice — tenants, billing, unified inbox.'),
        (12, 'API Reference', 'All 51 endpoints — public, admin, v1, webhooks, cron.'),
        (13, 'Background Jobs', '13 scheduled jobs orchestrating the lifecycle.'),
        (14, 'Data Dictionary', '88 tables, 197 RLS policies, 41 functions — fully indexed.'),
        (15, 'Integrations', '11 third-party services — Stripe, GHL, Twilio, Resend, IMAP, AI.'),
        (16, 'Security Audit', 'Authn, authz, secrets, RLS, compliance posture.'),
        (17, 'Operations & Resilience', 'Deploy, monitor, backup, incident response.'),
        (18, 'Technical Debt — by the numbers', 'Real counts. No hand-waving.'),
        (19, 'Recommendations & Roadmap', '30 / 90 / 180-day prioritized plan.'),
    ]
    add_toc(doc, toc_entries)

    # Load chapter modules
    chapters = [
        ('chapter_01_executive', 'render'),
        ('chapter_02_at_a_glance', 'render'),
        ('chapter_03_tech_stack', 'render'),
        ('chapter_04_multitenant', 'render'),
        ('chapter_05_routing', 'render'),
        ('chapter_06_public_site', 'render'),
        ('chapter_07_portal', 'render'),
        ('chapter_08_driver', 'render'),
        ('chapter_09_booking_wizard', 'render'),
        ('chapter_10_admin', 'render'),
        ('chapter_11_superadmin', 'render'),
        ('chapter_12_api', 'render'),
        ('chapter_13_crons', 'render'),
        ('chapter_14_schema', 'render'),
        ('chapter_15_integrations', 'render'),
        ('chapter_16_security', 'render'),
        ('chapter_17_operations', 'render'),
        ('chapter_18_tech_debt', 'render'),
        ('chapter_19_recommendations', 'render'),
    ]

    # Shared API surface for chapter modules
    api = {
        'doc': doc,
        'palette': {
            'NAVY': NAVY, 'NAVY_D': NAVY_D, 'GOLD': GOLD, 'INK': INK,
            'GRAY': GRAY, 'GRAY_L': GRAY_L, 'GREEN': GREEN, 'AMBER': AMBER,
            'RED': RED, 'BLUE': BLUE, 'TEAL': TEAL, 'PURPLE': PURPLE,
        },
        'add_h1': add_h1, 'add_h2': add_h2, 'add_h3': add_h3, 'add_h4': add_h4,
        'add_p': add_p, 'add_runs': add_runs,
        'add_bullet': add_bullet, 'add_numbered': add_numbered,
        'add_chapter_divider': add_chapter_divider,
        'add_section_break': add_section_break,
        'add_callout': add_callout,
        'add_kv_table': add_kv_table,
        'add_stats_strip': add_stats_strip,
        'add_data_sheet': add_data_sheet,
        'add_mono_block': add_mono_block,
        'add_pill_row': add_pill_row,
        'add_2col': add_2col,
    }

    for mod_name, fn_name in chapters:
        mod = safe_import(mod_name)
        if mod and hasattr(mod, fn_name):
            print(f'  rendering {mod_name}...')
            getattr(mod, fn_name)(api)
        else:
            # Placeholder if chapter file missing
            add_chapter_divider(doc, int(mod_name.split('_')[1]),
                                mod_name.replace('chapter_', '').replace('_', ' ').title(),
                                'Chapter under construction')

    out_path = os.path.join(REPO, 'RentalFlow_Technical_Audit.docx')
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f'\nSaved: {out_path} ({size_kb} KB)')
    return out_path


if __name__ == '__main__':
    main()
